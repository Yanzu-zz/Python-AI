# coding: utf-8

import os
from docx import Document
from docx.shared import Inches, RGBColor, Pt
from docx.enum.text import WD_PARAGRAPH_ALIGNMENT, WD_ALIGN_PARAGRAPH

doc = Document()

# 样式设置
style = doc.styles['Normal']

# 单个字体设置
style.font.name = '钉钉进步体'
style.font.color.rgb = RGBColor(255, 105, 0)
style.font.size = Pt(16)

# 添加标题，其中第二个参数是标题字体大小等级
# 这里一开始的 heading 无法设置斜体之类的
# 想要设置斜体可以一开始的内容为空，用追加的方式就可以设置斜体了
# title = doc.add_heading('My title', 0)
title = doc.add_heading('', 0)
# 标题设置居中
title.alignment = WD_PARAGRAPH_ALIGNMENT.CENTER
title.style.font.size = Pt(35)
# 追加内容（此时还是标题内容）
_t = title.add_run('My title\n1235')
# 设置斜体
_t.italic = True
_t.bold = True
_t.underline = True

# 添加段落
p = doc.add_paragraph('欢迎来到这里学习python word\n')
# 追加
p.add_run('这是关于 word 生成的基础知识').bold = True
# 这里不会设置成bold
p.add_run('后面追加的内容阿斯顿发生')
p.alignment = WD_PARAGRAPH_ALIGNMENT.RIGHT

# 这种方法可以控制居中
p1 = doc.add_paragraph()
p1.alingment = WD_ALIGN_PARAGRAPH.CENTER
_p1 = p1.add_run()
image1 = _p1.add_picture('./imgs/2022.jpg', width=Inches(2), height=Inches(2))

# 添加图片（这种无法控制位置）
image2 = doc.add_picture('./imgs/logo-2023.png', width=Inches(2), height=Inches(2))

# 添加表格
title = ['name', 'age', 'sex']
table = doc.add_table(rows=1, cols=3)
title_cells = table.rows[0].cells
title_cells[0].text = title[0]
title_cells[1].text = title[1]
title_cells[2].text = title[2]

data = [
    ('xiaomu', '10', 'man'),
    ('dewei', '34', 'man'),
    ('xiaoli', '18', 'woman')
]

# 标准就是用循环
for d in data:
    row_cells = table.add_row().cells
    for j in range(len(d)):
        row_cells[j].text = d[j]

# 分页功能（也就是生成新的一页）
doc.add_page_break()
# 给第二页加标题
title = doc.add_heading('New title2', 0)

doc.add_page_break()
title = doc.add_heading('My title3', 0)

# 生成 word
doc.save('write_word1.docx')
# 打开生成的 word
os.system('write_word1.docx')
