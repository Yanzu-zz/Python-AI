from docx import Document

import glob


class ReadDoc(object):
    def __init__(self, path):
        self.doc = Document(path)
        self.p_text = ''
        self.table_text = ''

        # 初始化自动获取文字和表格信息
        self.get_para()
        self.get_table()

    # 读取 word 文本，表格之类的数据
    def get_para(self):
        for p in self.doc.paragraphs:
            self.p_text += p.text + '\n'

    def get_table(self):
        for table in self.doc.tables:
            for row in table.rows:
                _cell_str = ''
                for cell in row.cells:
                    _cell_str += cell.text + ','
                self.table_text += _cell_str + '\n'


def search_word(path, targets):
    # 获取当前目录下所有文件的信息
    result = glob.glob(path)
    final_result = []

    for i in result:
        isuse = True
        if glob.os.path.isfile(i):
            # 确定是一个 .docx 文件才读取
            if i.endswith('.docx'):
                doc = ReadDoc(i)
                p_text = doc.p_text
                t_text = doc.table_text
                all_text = p_text + t_text

                # 满足所有条件的 docx 文档才是我们想要的
                # 比如文档里面 年龄<=35，技能要有C++等
                for target in targets:
                    if target not in all_text:
                        isuse = False
                        break
                if not isuse:
                    continue

                final_result.append(i)

    return final_result


if __name__ == '__main__':
    path = glob.os.path.join(glob.os.getcwd(), '*')
    res = search_word(path, ['C++','golang'])
    print(res)
