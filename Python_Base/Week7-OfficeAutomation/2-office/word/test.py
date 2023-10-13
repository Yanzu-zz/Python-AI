from docx import Document

doc = Document('./test.docx')

# print(doc.paragraphs)
# 打印文本
for p in doc.paragraphs:
    print(p.text)

# 打印表格
for t in doc.tables:
    # 遍历每行的表格
    for row in t.rows:
        # 获取每行的单个小表格
        _row = ''
        for cell in row.cells:
            _row += cell.text + ', '
        print(_row)
